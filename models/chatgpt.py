import os
import tempfile
from docx import Document
import openai
import concurrent.futures
import logging

# Logging setup
logging.basicConfig(level=logging.INFO, format='[%(threadName)s] %(message)s')
logger = logging.getLogger(__name__)

# OpenAI configuration
openai.organization = "org-2RniczpRsyTC1u0YUI5ebhAa"
openai.api_key = os.environ.get("OPENAI_KEY")

ALLOWED_EXTENSIONS_DOCX = {'docx'}

def allowed_file(filename, allowed_extensions=ALLOWED_EXTENSIONS_DOCX):
    logger.info(f"Checking if {filename} is an allowed file type.")
    return_value = '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions
    if return_value:
        logger.info(f"{filename} is an allowed file type.")
    else:
        logger.warning(f"{filename} is not an allowed file type.")
    return return_value

def process_file(file_path, prompt, progress_callback=None):
    logger.info(f"Starting processing for file: {file_path}")

    if not allowed_file(file_path):
        logger.error("Invalid file type.")
        return file_path, None

    try:
        return file_path, process_docx(file_path, prompt, progress_callback)
    except Exception as e:
        logger.error(f"Error processing the DOCX file: {str(e)}")
        return file_path, None

    logger.info(f"Finished processing for file: {file_path}")

def process_docx(file_path, prompt, progress_callback=None):
    if progress_callback:
        progress_callback("started")
    logger.info("Processing DOCX file.")
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        doc = Document(file_path)
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        CHUNK_SIZE = 4000
        text_chunks = [full_text[i:i + CHUNK_SIZE] for i in range(0, len(full_text), CHUNK_SIZE)]
        
        # Concurrently processing the text chunks
        logger.info(f"Starting concurrent processing of {len(text_chunks)} text chunks.")
        with concurrent.futures.ThreadPoolExecutor() as executor:
            results = list(executor.map(process_chunk, text_chunks, [prompt] * len(text_chunks)))

        base_name = os.path.splitext(os.path.basename(file_path))[0]
        safe_filename = ''.join([c if c.isalnum() else '_' for c in base_name])
        output_filename = f"{safe_filename}_processed.docx"
        output_path = os.path.join(tempfile.gettempdir(), output_filename)

        output_doc = Document()
        for result in results:
            output_doc.add_paragraph(result)

        output_doc.save(output_path)
        logger.info(f"Modified DOCX saved at {output_path}")
        if progress_callback:
            progress_callback("completed")
        
    logger.info(f"Finished processing DOCX: {output_path}")
    return output_path

def process_chunk(chunk, prompt):
    logger.info("Starting to process chunk.")
    try:
        messages = [{
            "role": "user",
            "content": f"{prompt} {chunk}"
        }]
        response = openai.ChatCompletion.create(model="gpt-3.5-turbo", messages=messages)
        return response.choices[0].message['content'].strip()
    except Exception as e:
        logger.error(f"Error processing chunk: {str(e)}")
        return ""
    logger.info("Finished processing chunk.")

if __name__ == "__main__":
    # Here, you can add a test function or a simple use-case to see the code in action.
    pass
